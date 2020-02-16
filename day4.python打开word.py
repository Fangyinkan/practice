# coding=utf-8
from docx import Document
import zipfile

# 读纯文本文档
# document = Document(r"C:\Users\KY\Desktop\python学习文件夹\长恨歌.docx")
# all_paragraphs = document.paragraphs
# for paragraph in all_paragraphs:
#     print(paragraph.text)
# 读表格类文档
# document = Document(r"C:\Users\KY\Desktop\python学习文件夹\长恨歌.docx")
# all_tables=document.tables
# for table in all_tables:
#     for row in table.rows:
#         for cell in row.cells:
#             print(cell.text)
# 读文本加表格类文档
word = zipfile.ZipFile(r"C:\Users\KY\Desktop\python学习文件夹\长恨歌.docx")  # 压缩word为zip
xml = word.read('word/document.xml').decode('utf-8')  # 读取word/document.xml文件
xml_list = xml.split('<w:t>')  # 将文件按照<w:t>切片 产生新列表
text_list = []  # 建一个空列表
for i in xml_list:
    if i.find('</w:t>') + 1:  # 如果找到</w:t>位置
        text_list.append(i[:i.find('</w:t>')])  # 添加从开始到</w:t>位置前的所有数据
    else:
        pass
text = "".join(text_list)  # 将列表内的元素合成一个字符串

print(text)
