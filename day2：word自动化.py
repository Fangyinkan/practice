# coding=utf-8
# 针对只改部分内容 不改大体内容 比如针对100个人发通知 抬头不同

from docx import Document  # 新建docx文件
from docx.enum.text import WD_ALIGN_PARAGRAPH  # 对齐方式
from docx.shared import Pt  # 磅数
from docx.oxml.ns import qn  # 中文格式
import time  # 导入时间

province_list = ['北京市', '天津市', '上海市', '重庆市', '河北省', '山西省', '辽宁省', '吉林省', '黑龙江省', '江苏省', '浙江省', '安徽省', '福建省', '江西省',
                 '山东省', '河南省', '湖北省', '湖南省', '广东省', '海南省', '四川省', '贵州省', '云南省', '陕西省', '甘肃省', '青海省', '台湾省', '内蒙古自治区',
                 '广西壮族自治区', '西藏自治区', '宁夏回族自治区', '新疆维吾尔自治区', '香港特别行政区', '澳门特别行政区']
today1 = time.strftime("%Y{y}%m{m}%d{d}", time.localtime()).format(y='年', m='月', d='日')  # 获取现在时间

for i in province_list:
    document = Document()
    document.styles['Normal'].font.name = u'宋体'  # 西文字体
    document.styles['Normal'].element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')  # 中文字体
    p1 = document.add_paragraph()  # 添加第一段
    p1.alignment = WD_ALIGN_PARAGRAPH.CENTER  # 中心对齐 默认左对齐
    run1 = p1.add_run('关于下达%s防疫的通知' % today1)  # 写内容
    run1.font.name = '微软雅黑'  # 内容西文字体
    run1.element.rPr.rFonts.set(qn('w:eastAsia'), u'微软雅黑')  # 内容中文字体
    run1.font.size = Pt(21)  # 字体大小
    run1.font.bold = True  # 字体加粗
    p1.space_after = Pt(5)  # 段后距离
    p1.space_before = Pt(5)  # 段前距离

    p2 = document.add_paragraph()
    run2 = p2.add_run('%s:' % i)
    run2.font.name = '仿宋GB_2312'
    run2.element.rPr.rFonts.set(qn('w:eastAsia'), u'仿宋GB_2312')
    run2.font.size = Pt(16)
    p2.space_after = Pt(5)
    p2.space_before = Pt(5)

    p3 = document.add_paragraph()
    run3 = p3.add_run('    经研究，现下达%s防疫通知给你们，请认真落实执行。' % today1)
    run3.font.name = '仿宋GB_2312'
    run3.element.rPr.rFonts.set(qn('w:eastAsia'), u'仿宋GB_2312')
    run3.font.size = Pt(16)
    p3.space_after = Pt(5)
    p3.space_before = Pt(5)

    p4 = document.add_paragraph()
    run4 = p4.add_run('国务院防疫小组')
    p4.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run4.font.name = '仿宋GB_2312'
    run4.element.rPr.rFonts.set(qn('w:eastAsia'), u'仿宋GB_2312')
    run4.font.size = Pt(16)
    p4.space_after = Pt(5)
    p4.space_before = Pt(5)

    p5 = document.add_paragraph()
    run5 = p5.add_run('%s' % today1)
    p5.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run5.font.name = '仿宋GB_2312'
    run5.element.rPr.rFonts.set(qn('w:eastAsia'), u'仿宋GB_2312')
    run5.font.size = Pt(16)
    p5.space_after = Pt(5)
    p5.space_before = Pt(5)

    document.save('%s通知.docx' % i)
